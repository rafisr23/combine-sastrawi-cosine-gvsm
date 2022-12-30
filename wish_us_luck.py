# Library
import sys
import fitz  # Membaca file PDF
import nltk  # Natural Language Toolkit
import re  # Menghapus karakter angka.
import string  # Menghapus karakter tanda baca.
import matplotlib.pyplot as plt  # Menggambarkan Frekuensi Kemunculan
import docx  # Membaca File docx
# from docx2pdf import convert  # Meng-convert docx ke pdf
import numpy as np
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory  # Stemming Sastrawi
from Sastrawi.StopWordRemover.StopWordRemoverFactory import StopWordRemoverFactory, StopWordRemover, ArrayDictionary  # Stopword Sastrawi
import os  # Membaca file di dalam folder
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2

# library pyqt5 untuk GUI
from PyQt5 import QtCore, QtWidgets
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUi

class ShowGUI(QMainWindow):
    def __init__(self):
        super(ShowGUI, self).__init__()
        loadUi('gui.ui', self)
        self.file = ""
        self.query = ""
        self.actionOpen.triggered.connect(self.openClicked)
        self.buttonGvsm.clicked.connect(self.main)
        self.buttonQuery.clicked.connect(self.insertQuery)

    # Get Dir
    def openClicked(self):
        self.file = str(QFileDialog.getExistingDirectory(self, "Select Directory"))
        self.pathLabel.setText(str(self.file))
        count = 0
        # Iterate directory
        for path in os.listdir(self.file):
            # check if current path is a file
            if os.path.isfile(os.path.join(self.file, path)):
                count += 1
        self.totalFileLabel.setText(str(count))

    def insertQuery(self):
        self.query = self.queryLabel.toPlainText()
        print(self.query)

    # Input File - DOCX
    def getDOCX(self, filename):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)

    # Input File - PDF
    def getPDF(self, filename):
        with fitz.open(filename) as doc:  # Membuka File PDF
            kalimat = ""
            for page in doc:
                kalimat += page.get_text()  # Memasukkan Kata ke variabel Kalimat
        return kalimat

    # Membuat Value String Menjadi List
    def convert(self, value):
        li = list(value.split(" "))
        return li

    # Memproses File Inputan
    def proses(self, file):
        doc = file
        if doc.endswith('.pdf'):
            hasil = self.getPDF(doc)
            hasil_stemming = self.getStemming(hasil)
            hasil_fix = self.convert(hasil_stemming)
            kemunculan = nltk.FreqDist(self.convert(hasil_stemming))
            # print('\nStemming hasil -> ' + str(doc) + ':')
            # print(kemunculan.most_common())

        elif doc.endswith('.docx'):
            hasil = self.getDOCX(doc)
            hasil_stemming = self.getStemming(hasil)
            hasil_fix = self.convert(hasil_stemming)
            kemunculan = nltk.FreqDist(self.convert(hasil_stemming))
            # print('\nStemming hasil -> ' + str(doc) + ':')
            # print(kemunculan.most_common())

        else:
            print('Harap Masukkan File PDF atau DOCX')

        return hasil_fix

    def retrieve_relevant_documents(self, folder_path, query):
        # Create a stemmer object
        factory = StemmerFactory()
        stemmer = factory.create_stemmer()

        # Read the documents from the specified folder
        documents = []
        filenames = []
        for filename in os.listdir(folder_path):
            if filename.endswith(".pdf"):
                filenames.append(filename)
                with open(os.path.join(folder_path, filename), "rb") as f:
                    pdf_reader = PyPDF2.PdfFileReader(f)
                    num_pages = pdf_reader.getNumPages()
                    document = ""
                    for page_num in range(num_pages):
                        page = pdf_reader.getPage(page_num)
                        document += page.extractText()
                    documents.append(document)

            elif filename.endswith(".docx"):
                filenames.append(filename)
                doc = docx.Document(os.path.join(folder_path, filename))
                document = ""
                for paragraph in doc.paragraphs:
                    document += paragraph.text
                documents.append(document)

        # Preprocess the documents by stemming the words
        processed_documents = []
        for document in documents:
            processed_document = " ".join([stemmer.stem(word) for word in document.split()])
            processed_documents.append(processed_document)

        # Create a vocabulary of all the unique terms in the documents
        vocab = set()
        for document in processed_documents:
            for word in document.split():
                vocab.add(word)

        vocab = list(vocab)

        # Create a matrix of document vectors, with each row representing a document and each column representing a term
        matrix = np.zeros((len(processed_documents), len(vocab)))

        # Calculate the term frequency of each term in each document
        for i, document in enumerate(processed_documents):
            for j, term in enumerate(vocab):
                matrix[i, j] = document.split().count(term)

        # Perform term weighting using the GVSM
        matrix = matrix / np.sqrt(np.sum(matrix ** 2, axis=1, keepdims=True))

        # Preprocess the query by stemming the words
        processed_query = " ".join([stemmer.stem(word) for word in query.split()])

        # Create a query vector and perform term weighting using the GVSM
        query_vector = np.zeros(len(vocab))
        for j, term in enumerate(vocab):
            query_vector[j] = processed_query.split().count(term)
        query_vector = query_vector / np.sqrt(np.sum(query_vector ** 2))

        # Calculate the cosine similarity between the query vector and each document vector
        similarities = []
        for i, document_vector in enumerate(matrix):
            dot_product = np.dot(query_vector, document_vector)
            query_length = np.sqrt(np.dot(query_vector, query_vector))
            doc_length = np.sqrt(np.dot(document_vector, document_vector))
            similarity = dot_product / (query_length * doc_length)
            similarities.append(similarity)

        # Rank the documents by similarity and retrieve the most relevant ones
        ranked_documents = sorted(zip(similarities, documents, filenames), reverse=True)
        return ranked_documents
        # Rank the documents by similarity and retrieve the most relevant ones
        # ranked_documents = sorted(zip(similarities, documents), reverse=True)
        # filenames = [document.replace(folder_path + "/", "") for _, document in ranked_documents]
        # return filenames[:2]

    # Retrieve the most relevant documents
    def main(self):
        # Define the folder path containing the documents
        folder_path = self.file

        # Define the query
        query = self.query
        most_relevant_filenames = self.retrieve_relevant_documents(folder_path, query)

        print("Most relevant documents:")
        for similarity, document, filename in most_relevant_filenames:
            # filename = document.replace(folder_path + "/", "")
            similarity_percentage = round(similarity * 100)
            similarity = similarity
            print(f"Document: {filename} (similarity: {similarity}) ({similarity_percentage}%)")
        # for filename in most_relevant_filenames:
        #     print(f"Filename: {filename}")

app = QtWidgets.QApplication(sys.argv)
window = ShowGUI()
window.setWindowTitle('Information Retrieval')
window.show()
sys.exit(app.exec_())
